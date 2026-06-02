#!/usr/bin/env python3
"""Build the GFM Missouri Disability + IME outreach packet (Word)."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

NAVY = RGBColor(0x1F, 0x38, 0x64)
ORANGE = RGBColor(0xC5, 0x5A, 0x11)
GREY = RGBColor(0x59, 0x59, 0x59)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = NAVY
    return p

def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = ORANGE
    p.space_before = Pt(8)
    return p

def small(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(9); r.font.color.rgb = GREY
    return p

def body(text):
    return doc.add_paragraph(text)

def rule():
    p = doc.add_paragraph()
    r = p.add_run("—" * 38); r.font.color.rgb = RGBColor(0xBF,0xBF,0xBF)

# ---------- Cover / header ----------
h1("Garcia Family Medicine")
p = doc.add_paragraph()
r = p.add_run("Disability Evaluations & Independent Medical Evaluations — Outreach Packet")
r.bold = True; r.font.color.rgb = NAVY; r.font.size = Pt(12)
small("Dr. Tess Garcia, MD • Dual board-certified • 26+ years • 801 NW St. Mary's Dr., Suite 209, Blue Springs, MO 64014 • (816) 330-7575")
small("Outreach: Gigi Escalante Garcia • IME line: veteransIME@garciafamilymedicine.care • garciafamilymedicine.care")
rule()

h2("How to use this packet")
body("Five ready-to-send templates, one per channel. Fill the [brackets], attach the items listed under “Provider credentialing — have these ready,” and send. Pair each send with the matching row(s) in the Missouri Disability + IME list. Verify the current recipient/contact for each organization before sending (provider-intake contacts change).")

# ---------- Credentialing checklist ----------
h2("Provider credentialing — have these ready (most networks ask for these)")
for item in [
    "Dr. Tess Garcia CV (with board certifications and 26+ years experience highlighted)",
    "Active Missouri medical license + license number; DEA if requested",
    "NPI number; CAQH ProView profile (keep it attested/current)",
    "Board certification certificates",
    "Malpractice / professional liability declarations page (coverage limits)",
    "W-9 for Garcia Family Medicine; practice address, hours, languages, accessibility",
    "Sample work product: the completed IME report (e.g., M. Mendez) and a Nexus letter sample (redacted)",
    "IME fee schedule: $3,000 retainer (+$300/hr for extended review); Nexus letters $1,500–$3,000",
]:
    doc.add_paragraph(item, style="List Bullet")

rule()

# ---------- Template A: VA C&P TPAs ----------
h2("Template A — VA disability (C&P) exam TPAs (QTC, VES, Optum Serve, Loyal Source, Sterling)")
small("Goal: get credentialed as a network examiner for veteran Compensation & Pension disability exams.")
body("Subject: Missouri provider application — Garcia Family Medicine (C&P / disability exams)")
body("Hello [Provider Network / Credentialing Team],")
body("I'm writing on behalf of Garcia Family Medicine in Blue Springs, Missouri. Dr. Tess Garcia is a dual board-certified physician with 26+ years of clinical experience, and we are interested in joining your examiner network to perform Compensation & Pension / disability evaluations for veterans in the Kansas City metro and surrounding Missouri counties.")
body("We offer fast, thorough, defensible reports and have current capacity for in-person and telehealth evaluations. Could you send the provider application and credentialing requirements, and let us know the current exam volume/need in our area?")
body("Dr. Tess holds an active Missouri license and board certifications; CV, NPI, malpractice, and CAQH are ready on request. Thank you for your time — we'd be glad to get started.")
body("Warm regards,\n[Name], [Title]\nGarcia Family Medicine • (816) 330-7575 • veteransIME@garciafamilymedicine.care")

rule()

# ---------- Template B: Missouri DDS (SSA CE) ----------
h2("Template B — Social Security disability: Missouri DDS Consultative Exam (CE) provider")
small("Goal: enroll to perform SSA Consultative Exams for Missouri claimants. Ask for the Professional Relations Officer.")
body("Subject: Consultative Examination (CE) provider enrollment — Garcia Family Medicine, Blue Springs MO")
body("Hello,")
body("I'd like to inquire about enrolling Garcia Family Medicine as a Consultative Examination provider for Missouri Disability Determination Services. Dr. Tess Garcia is a dual board-certified physician with 26+ years of experience, located in Blue Springs (Jackson County), with capacity for adult physical and mental status consultative exams.")
body("Could you connect me with the Professional Relations Officer (or the correct contact) and share the CE provider application, fee schedule, and documentation requirements? We can offer prompt scheduling and timely, complete reports.")
body("Thank you,\n[Name], [Title]\nGarcia Family Medicine • (816) 330-7575")

rule()

# ---------- Template C: IME networks ----------
h2("Template C — IME networks (ExamWorks, MES, Dane Street, MCN, Genex, IMX, CorVel)")
small("Goal: join the physician panel for workers' comp / auto / liability IMEs in Missouri.")
body("Subject: Physician panel application — Missouri IME examiner (Garcia Family Medicine)")
body("Hello [Panel / Provider Recruitment Team],")
body("Garcia Family Medicine (Blue Springs, MO) would like to join your physician panel for Independent Medical Evaluations. Dr. Tess Garcia is dual board-certified with 26+ years of experience and performs workers' compensation, auto, and liability IMEs, record reviews, and impairment-related evaluations for Missouri cases.")
body("We provide clear, defensible reports with fast turnaround (telehealth and in-person available). Please send your panel application and credentialing checklist, and let us know current case volume in the Kansas City / western Missouri area.")
body("Best regards,\n[Name], [Title]\nGarcia Family Medicine • (816) 330-7575 • garciafamilymedicine.care")

rule()

# ---------- Template D: Attorneys ----------
h2("Template D — Workers' comp & personal-injury attorneys (e.g., Boyd, Kenter, Thomas & Parrish; MATA members)")
small("Goal: become the go-to IME / medical-opinion expert. Warm for BKTP — you've already delivered a report for them.")
body("Subject: Independent Medical Evaluations & medical opinion letters — Garcia Family Medicine")
body("Dear [Attorney / Firm],")
body("Thank you for the opportunity to support [the Mendez matter / your clients]. I wanted to make sure your team knows Garcia Family Medicine offers Independent Medical Evaluations, record reviews, and physician-signed medical opinion letters for workers' compensation and personal-injury cases.")
body("Dr. Tess Garcia is dual board-certified with 26+ years of experience and delivers concise, defensible reports with responsive turnaround. Our IME engagement begins with a $3,000 retainer (plus $300/hour only if extended record review is required), and we're always transparent about scope before proceeding.")
body("I've attached our IME overview and a work sample. We'd welcome the chance to be a resource for your next case — may I send our intake form and current availability?")
body("With appreciation,\nGigi Escalante Garcia, Outreach Director\nGarcia Family Medicine • (816) 330-7575 • veteransIME@garciafamilymedicine.care")

rule()

# ---------- Template E: Veteran orgs / VSOs ----------
h2("Template E — Veteran orgs & County Veterans Service Officers (Missouri Veterans Commission, DAV, VFW, American Legion, CVSOs)")
small("Goal: become the trusted local source for Nexus letters and Veterans@Work IMEs.")
body("Subject: Partnering to support Missouri veterans — Nexus letters & independent medical evaluations")
body("Hello [Service Officer / Department Contact],")
body("Garcia Family Medicine in Blue Springs supports veterans through our Veterans@Work IME pathway and physician-prepared Nexus letters that connect a veteran's service to their current condition. Dr. Tess Garcia is dual board-certified with 26+ years of experience and brings a faith-and-family, veteran-centered approach.")
body("Many veterans you assist need strong, well-documented medical evidence for their claims. We'd value the chance to be a referral resource for your service officers — with clear pricing, fast turnaround (72-hour options), and compassionate care. Could we schedule a brief call, or may I send referral flyers and a sample MoU for your team?")
body("Respectfully,\nGigi Escalante Garcia, Outreach Director\nGarcia Family Medicine • (816) 330-7575 • veteransIME@garciafamilymedicine.care")

rule()
small("Note: All fees and turnaround times above reflect your existing IME/Nexus service documents. Confirm the current provider-intake contact for each organization (see the Missouri Disability + IME list) before sending. Dr. Tess does not promise claim approval; reports are unbiased and evidence-based.")

out = "/home/user/voicerx-site/GFM_Missouri_Disability_IME_Outreach_Packet.docx"
doc.save(out)
print("saved", out)
